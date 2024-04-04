import pandas as pd
import matplotlib.pyplot as plt
import os
from docx import Document

# 读取CSV文件
excel_file = pd.ExcelFile('D:\\Desktop\\BUPT\\Final Project\\samples\\DescargasDatosUCIN.xlsx', engine='openpyxl')

df = pd.DataFrame()
for sheet_name in excel_file.sheet_names:
    sheet_df = excel_file.parse(sheet_name)
    df = pd.concat([df, sheet_df], ignore_index=True)
# 确保时间戳列是datetime类型
df['horaRegistro'] = pd.to_datetime(df['horaRegistro'])

# 获取所有不同的数据类型和病床号
data_types = df['dato'].unique()
beds = df['cama'].unique()
bed_count_per_data_type = df.groupby('dato')['cama'].nunique()

# 输出每种数据类型对应的床位数量
print("每种数据类型对应的床位数量:")
print(bed_count_per_data_type)
doc = Document()
# 添加标题
doc.add_heading('每种数据类型对应的床位数量', level=1)

# 添加表格
table = doc.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '数据类型'
hdr_cells[1].text = '床位数量'

# 将数据添加到表格中
for data_type, bed_count in bed_count_per_data_type.items():
    row_cells = table.add_row().cells
    row_cells[0].text = str(data_type)
    row_cells[1].text = str(bed_count)

# 保存Word文档
doc.save('bed_count_per_data_type.docx')
doc.save('D:\\Desktop\\BUPT\\Final Project\\data_Analysis.docx')

# 循环遍历所有数据类型和病床号，并生成对应的图像并保存到电脑上
# for data_type in data_types:
#     for bed in beds:
#         # 筛选特定床位和数据类型的数据
#         filtered_df = df[(df['cama'] == bed) & (df['dato'] == data_type)]
#
#         # 对数据按时间排序
#         filtered_df.sort_values('horaRegistro', inplace=True)
#
#         # 绘制数据变化图
#         plt.figure(figsize=(10, 6))
#         plt.plot(filtered_df['horaRegistro'], filtered_df['valor'], marker='o', linestyle='-')
#         plt.title(f'Changes in {data_type} data overtime for bed {bed} ')
#         plt.xlabel('time')
#         plt.ylabel('value')
#         plt.xticks(rotation=45)
#         plt.tight_layout()
#
#         # 创建文件夹（如果不存在）
#         output_dir = 'D:\\Desktop\\BUPT\\Final Project\\plots'
#         if not os.path.exists(output_dir):
#             os.makedirs(output_dir)
#
#         # 保存图像到电脑上
#         file_name = f'{output_dir}\\{data_type}_{bed}.png'
#         plt.savefig(file_name)
#
#         # 关闭图形，以便释放资源
#         plt.close()

print("图像已保存到电脑上。")
