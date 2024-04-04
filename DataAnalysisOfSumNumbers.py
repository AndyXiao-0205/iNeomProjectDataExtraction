import pandas as pd
import matplotlib.pyplot as plt
from docx import Document

df = pd.read_csv('D:\\Desktop\\BUPT\\Final Project\\DescargasDatosUCIN.csv')
data_counts = df['dato'].value_counts()

data_ranges = df.groupby('dato')['valor'].agg(['min', 'max'])
data_type_counts = df['dato'].value_counts()
# 输出各种数据类型的范围（最小值和最大值）
print("各种数据类型的范围:")
print(data_ranges)

doc = Document()

# 添加标题
doc.add_heading('Ranges of data', level=1)

# 添加表格
table1 = doc.add_table(rows=1, cols=3)
hdr_cells = table1.rows[0].cells
hdr_cells[0].text = 'Data type'
hdr_cells[1].text = 'Min'
hdr_cells[2].text = 'Max'

# 添加标题
doc.add_heading('Numbers of data', level=1)


# 添加表格
table2 = doc.add_table(rows=1, cols=2)
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Data type'
hdr_cells[1].text = 'Number'

# 将数据添加到表格中
for index, count in data_type_counts.items():
    row_cells = table2.add_row().cells
    row_cells[0].text = str(index)
    row_cells[1].text = str(count)

# 将数据添加到表格中
for index, row in data_ranges.iterrows():
    row_cells = table1.add_row().cells
    row_cells[0].text = str(index)
    row_cells[1].text = str(row['min'])
    row_cells[2].text = str(row['max'])

# 保存Word文档
doc.save('D:\\Desktop\\BUPT\\Final Project\\data_Analysis.docx')

# print(data_counts)


# data_counts.plot(kind='bar')
# plt.title('Numbers of Datas')
# plt.xlabel('Dato(Data type)')
# plt.ylabel('Numbers')
# plt.xticks(rotation=45)
# plt.tight_layout()
# plt.show()

# # 创建包含数据类型和对应行数的DataFrame
# data_counts_df = pd.DataFrame(data_counts)
# data_counts_df.columns = ['行数']

# # 输出总共有多少种数据
# total_data_types = len(data_counts_df)
# print(f"总共有 {total_data_types} 种数据类型.")

# # 输出包含数据类型和对应行数的表格
# print("各种数据类型及其行数:")
# print(data_counts_df)